from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from vector_store import VectorStore
from config import Config
import pypdf
from docx import Document as DocxDocument
import re
from embedding_cache import EmbeddingCache
from logging_config import rag_logger


class AdvancedDocumentProcessor:
    """Advanced document processor tailored for French academic PDFs with numbered sections"""
    
    # Regex for French academic document headers
    HEADER_REGEX = re.compile(
        r'^(Chapitre\s+\d+|'                    # Chapitre 3
        r'[IVX]+\.\s+|'                         # I. II. III. IV. V. (Roman numerals)
        r'\d+(\.\d+)*\.?\s+)',                  # 5.1.3.2. Title or 1. Title
        re.IGNORECASE | re.MULTILINE
    )
    
    # Noise patterns to filter out
    NOISE_PATTERNS = [
        r'Page\s+\d+/\d+',
        r'Cours\s+Outils',
        r'N\.BOHLI',
        r'IIA4',
        r'^\s*$'  # Empty lines
    ]
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 120):
        """
        Initialize with settings optimized for academic PDFs
        
        Args:
            chunk_size: Target size (800 tokens ≈ good for technical content)
            chunk_overlap: Overlap (15% recommended)
        """
        self.config = Config()
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = rag_logger
        
        # Initialize embedding cache
        self.embedding_cache = EmbeddingCache()
        
        # Text splitter for long sections
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
            keep_separator=True  # Keep context
        )
        
        print(f" Chunking Strategy (Academic PDF Optimized):")
        print(f"   - Section-based chunking with numbered headers")
        print(f"   - Chunk size: {chunk_size} chars (~{chunk_size//4} tokens)")
        print(f"   - Overlap: {chunk_overlap} chars ({int(chunk_overlap/chunk_size*100)}%)")
        print(f"   - Contextual headers: Enabled")
        print(f"   - Noise filtering: Enabled")
        print(f"   - Embedding cache: Enabled")
    
    def is_noise(self, line: str) -> bool:
        """Check if a line is PDF noise (headers, footers, etc.)"""
        for pattern in self.NOISE_PATTERNS:
            if re.search(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_section_header(self, line: str) -> bool:
        """Detect if line is a section header (Chapitre X, numbered sections)"""
        return bool(self.HEADER_REGEX.match(line.strip()))
    
    def extract_sections_from_pdf(self, pdf_path: Path) -> List[Dict[str, Any]]:
        """
        Extract structured sections from PDF based on numbered headers
        
        Returns list of sections with:
        - chapter: Chapitre number/title
        - section: Section number/title
        - page_num: Page number
        - text: Content
        """
        try:
            reader = pypdf.PdfReader(str(pdf_path))
            
            # Extract all text first
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"
            
            # Split by headers using more aggressive regex
            # Matches: "I. Title", "II. Title", "1. Title", "2.3. Title", "Chapitre X"
            header_pattern = re.compile(
                r'((?:^|\s+)(Chapitre\s+\d+[^.]*|[IVX]+\.\s+[A-ZÀ-Ü][^.]{3,}|'
                r'\d+(?:\.\d+)*\.\s+[A-ZÀ-Ü][^.]{3,}))',
                re.MULTILINE
            )
            
            # Find all headers with their positions
            headers_found = []
            for match in header_pattern.finditer(full_text):
                header_text = match.group(0).strip()
                # Skip noise
                if self.is_noise(header_text):
                    continue
                headers_found.append({
                    'text': header_text,
                    'start': match.start(),
                    'end': match.end()
                })
            
            # Extract sections between headers
            sections = []
            current_chapter = "Unknown Chapter"
            
            for i, header in enumerate(headers_found):
                # Get content from this header to next header (or end)
                start_pos = header['end']
                end_pos = headers_found[i+1]['start'] if i+1 < len(headers_found) else len(full_text)
                content = full_text[start_pos:end_pos].strip()
                
                # Check if this is a chapter header
                if header['text'].startswith("Chapitre"):
                    current_chapter = header['text']
                    section_title = ""
                else:
                    section_title = header['text']
                
                # Clean content
                clean_lines = []
                for line in content.split('\n'):
                    line = line.strip()
                    if line and not self.is_noise(line):
                        clean_lines.append(line)
                
                clean_content = '\n'.join(clean_lines)
                
                if clean_content:
                    sections.append({
                        "chapter": current_chapter,
                        "section": section_title,
                        "page_num": 1,
                        "text": clean_content
                    })
            
            return sections
            
        except Exception as e:
            print(f" Error extracting PDF sections: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def create_contextual_header(self, section_info: Dict[str, Any], file_name: str) -> str:
        """Create rich contextual header for embeddings"""
        parts = [f"Document: {file_name}"]
        
        if section_info.get("chapter"):
            parts.append(f"Chapitre: {section_info['chapter']}")
        
        if section_info.get("section"):
            parts.append(f"Section: {section_info['section']}")
        
        if section_info.get("page_num"):
            parts.append(f"Page: {section_info['page_num']}")
        
        header = " | ".join(parts)
        return f"[CONTEXT: {header}]\n\n"
    
    def chunk_section_with_overlap(self, section_info: Dict[str, Any], file_name: str) -> List[Document]:
        """
        Chunk a single section with sliding windows (for long sections)
        
        Strategy:
        - If section < chunk_size: return as-is
        - If section > chunk_size: split with overlap
        - Always preserve contextual headers
        """
        section_text = section_info["text"]
        
        # Create contextual header
        context_header = self.create_contextual_header(section_info, file_name)
        
        # Base metadata
        base_metadata = {
            "file_name": file_name,
            "chapter": section_info.get("chapter", "Unknown"),
            "section": section_info.get("section", "Unknown"),
            "page_num": section_info.get("page_num", 0)
        }
        
        # If section is small enough, return as single chunk
        if len(section_text) <= self.chunk_size:
            full_text = context_header + section_text
            return [Document(page_content=full_text, metadata=base_metadata)]
        
        # Otherwise, split with overlap
        text_with_context = context_header + section_text
        
        chunks = self.text_splitter.create_documents(
            texts=[text_with_context],
            metadatas=[base_metadata]
        )
        
        # Add sub-chunk indices
        for idx, chunk in enumerate(chunks):
            chunk.metadata["sub_chunk"] = idx
            chunk.metadata["total_sub_chunks"] = len(chunks)
        
        return chunks
    
    def process_pdf(self, pdf_path: Path) -> List[Document]:
        """Process PDF with section-based + sliding window chunking"""
        print(f"\n Processing PDF: {pdf_path.name}")
        
        # Check embedding cache first
        cached_data = self.embedding_cache.get(str(pdf_path))
        if cached_data:
            print(f"    [CACHE HIT] Loading from cache (FAST)")
            self.logger.info(f"Embedding cache hit for {pdf_path.name}")
            # Reconstruct Document objects from cached data
            all_chunks = []
            for chunk_data in cached_data['embeddings']:
                doc = Document(
                    page_content=chunk_data['content'],
                    metadata=chunk_data['metadata']
                )
                all_chunks.append(doc)
            print(f"    Loaded {len(all_chunks)} chunks from cache")
            return all_chunks
        
        print(f"    [CACHE MISS] Processing document...")
        self.logger.info(f"Processing {pdf_path.name} (cache miss)")
        
        # Extract sections
        sections = self.extract_sections_from_pdf(pdf_path)
        
        if not sections:
            print(f"    No sections found in {pdf_path.name}")
            return []
        
        print(f"    Found {len(sections)} sections")
        
        # Chunk each section
        all_chunks = []
        for section_info in sections:
            chunks = self.chunk_section_with_overlap(section_info, pdf_path.name)
            all_chunks.extend(chunks)
        
        print(f"    Created {len(all_chunks)} chunks ({len(all_chunks) - len(sections)} overlapping)")
        
        # Add global chunk indices
        for idx, chunk in enumerate(all_chunks):
            chunk.metadata["chunk_index"] = idx
        
        # Cache the chunks for future use
        chunk_data_list = []
        for chunk in all_chunks:
            chunk_data_list.append({
                'content': chunk.page_content,
                'metadata': chunk.metadata
            })
        
        self.embedding_cache.set(
            str(pdf_path),
            chunk_data_list,
            metadata={
                'num_chunks': len(all_chunks),
                'num_sections': len(sections),
                'file_name': pdf_path.name
            }
        )
        print(f"    Cached chunks for future use")
        self.logger.info(f"Cached {len(all_chunks)} chunks for {pdf_path.name}")
        
        return all_chunks
    
    def process_docx(self, docx_path: Path) -> List[Document]:
        """Process DOCX (simpler structure)"""
        print(f"\n Processing DOCX: {docx_path.name}")
        
        # Check embedding cache first
        cached_data = self.embedding_cache.get(str(docx_path))
        if cached_data:
            print(f"    [CACHE HIT] Loading from cache (FAST)")
            self.logger.info(f"Embedding cache hit for {docx_path.name}")
            all_chunks = []
            for chunk_data in cached_data['embeddings']:
                doc = Document(
                    page_content=chunk_data['content'],
                    metadata=chunk_data['metadata']
                )
                all_chunks.append(doc)
            print(f"    Loaded {len(all_chunks)} chunks from cache")
            return all_chunks
        
        print(f"    [CACHE MISS] Processing document...")
        self.logger.info(f"Processing {docx_path.name} (cache miss)")
        
        try:
            doc = DocxDocument(str(docx_path))
            all_chunks = []
            
            current_section = {"header": "Document Start", "paragraphs": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text or self.is_noise(text):
                    continue
                
                # Detect headers
                if para.style.name.startswith('Heading') or self.is_section_header(text):
                    if current_section["paragraphs"]:
                        # Process previous section
                        section_text = "\n".join(current_section["paragraphs"])
                        context = f"[CONTEXT: Document: {docx_path.name} | Section: {current_section['header']}]\n\n"
                        
                        chunks = self.text_splitter.create_documents(
                            texts=[context + section_text],
                            metadatas=[{
                                "file_name": docx_path.name,
                                "section": current_section['header']
                            }]
                        )
                        all_chunks.extend(chunks)
                    
                    current_section = {"header": text, "paragraphs": []}
                else:
                    current_section["paragraphs"].append(text)
            
            # Last section
            if current_section["paragraphs"]:
                section_text = "\n".join(current_section["paragraphs"])
                context = f"[CONTEXT: Document: {docx_path.name} | Section: {current_section['header']}]\n\n"
                chunks = self.text_splitter.create_documents(
                    texts=[context + section_text],
                    metadatas=[{"file_name": docx_path.name, "section": current_section['header']}]
                )
                all_chunks.extend(chunks)
            
            print(f"    Created {len(all_chunks)} chunks")
            
            # Add indices
            for idx, chunk in enumerate(all_chunks):
                chunk.metadata["chunk_index"] = idx
            
            # Cache the chunks
            chunk_data_list = []
            for chunk in all_chunks:
                chunk_data_list.append({
                    'content': chunk.page_content,
                    'metadata': chunk.metadata
                })
            
            self.embedding_cache.set(
                str(docx_path),
                chunk_data_list,
                metadata={
                    'num_chunks': len(all_chunks),
                    'file_name': docx_path.name
                }
            )
            print(f"    Cached chunks for future use")
            self.logger.info(f"Cached {len(all_chunks)} chunks for {docx_path.name}")
            
            return all_chunks
            
        except Exception as e:
            print(f" Error processing DOCX: {e}")
            self.logger.error(f"Error processing DOCX {docx_path.name}: {e}")
            return []
    
    def process_directory(self, directory_path: str = None) -> List[Document]:
        """Process all documents"""
        docs_path = Path(directory_path or self.config.DOCUMENTS_PATH)
        
        if not docs_path.exists():
            print(f" Directory {docs_path} does not exist!")
            return []
        
        print(f"\nProcessing documents from {docs_path}...")
        print(f"Strategy: Section-based + Contextual Headers + 15% Overlap\n")
        
        all_documents = []
        
        # Process PDFs
        for pdf_file in docs_path.glob("*.pdf"):
            chunks = self.process_pdf(pdf_file)
            all_documents.extend(chunks)
        
        # Process DOCX
        for docx_file in docs_path.glob("*.docx"):
            chunks = self.process_docx(docx_file)
            all_documents.extend(chunks)
    
        print(f"\nTotal chunks created: {len(all_documents)}")
        print(f"Quality improvements:")
        print(f"   - Noise filtered (Page X/Y, headers)")
        print(f"   - Section-aware chunking")
        print(f"   - Contextual headers added")
        print(f"   - Overlapping windows for continuity")
        
        # Display cache statistics
        cache_stats = self.embedding_cache.get_stats()
        print(f"\nEmbedding Cache Statistics:")
        print(f"   - Cache hits: {cache_stats['hits']}")
        print(f"   - Cache misses: {cache_stats['misses']}")
        print(f"   - Hit rate: {cache_stats['hit_rate']}")
        print(f"   - Cached files: {cache_stats['cached_files']}")
        print(f"   - Cache size: {cache_stats['cache_size_mb']} MB")
        self.logger.info(f"Document processing complete. Cache stats: {cache_stats}")
        
        return all_documents
    
    def build_vector_store(self, documents: List[Document], use_local: bool = True):
        """Build vector store with embedding caching for performance"""
        print(f"\nBuilding vector store with {len(documents)} chunks...")
        
        vector_store = VectorStore()
        
        if not vector_store.initialize(use_local_embeddings=use_local):
            print("ERROR: Failed to initialize vector store")
            return False
        
        try:
            # Group documents by file to leverage embedding cache
            files_docs = {}
            for doc in documents:
                file_name = doc.metadata.get('file_name', 'unknown')
                if file_name not in files_docs:
                    files_docs[file_name] = []
                files_docs[file_name].append(doc)
            
            print(f"\nProcessing embeddings for {len(files_docs)} files...")
            
            total_cached = 0
            total_computed = 0
            
            for file_name, file_docs in files_docs.items():
                # Create cache key based on document contents (not file hash)
                import hashlib
                content_hash = hashlib.md5("".join([d.page_content for d in file_docs]).encode()).hexdigest()
                cache_key = f"embeddings_{file_name}_{content_hash}"
                
                # Try to get cached embeddings
                # We store as a fake "file" in cache - using content instead of actual file
                cache_path = self.embedding_cache.cache_dir / f"{content_hash}.pkl"
                cached_embeddings = None
                if cache_path.exists():
                    import pickle
                    with open(cache_path, 'rb') as f:
                        cached_embeddings = pickle.load(f)
                
                if cached_embeddings:
                    # Use cached embeddings
                    print(f"   [{file_name}] Using cached embeddings ({len(file_docs)} chunks)")
                    self.logger.info(f"Using cached embeddings for {file_name}")
                    
                    embeddings = cached_embeddings['embeddings']
                    vector_store.add_documents_with_embeddings(file_docs, embeddings)
                    total_cached += len(file_docs)
                else:
                    # Compute embeddings
                    print(f"   [{file_name}] Computing embeddings ({len(file_docs)} chunks)...")
                    self.logger.info(f"Computing embeddings for {file_name}")
                    
                    # Get texts to embed
                    texts = [doc.page_content for doc in file_docs]
                    
                    # Compute embeddings using the embedding model
                    embeddings = vector_store.embedding_model.embed_documents(texts)
                    
                    # Cache the embeddings
                    import pickle
                    from datetime import datetime
                    cache_data = {
                        'embeddings': embeddings,
                        'metadata': {
                            'file_name': file_name,
                            'num_chunks': len(file_docs),
                            'embedding_dim': len(embeddings[0]) if embeddings else 0,
                            'timestamp': datetime.now().isoformat()
                        }
                    }
                    with open(cache_path, 'wb') as f:
                        pickle.dump(cache_data, f)
                    self.logger.info(f"Cached embeddings for {file_name}")
                    
                    # Add to vector store with pre-computed embeddings
                    vector_store.add_documents_with_embeddings(file_docs, embeddings)
                    total_computed += len(file_docs)
            
            # Summary
            print(f"\nEmbedding Processing Summary:")
            print(f"   - Cached: {total_cached} chunks")
            print(f"   - Computed: {total_computed} chunks")
            print(f"   - Total: {total_cached + total_computed} chunks")
            
            if total_cached + total_computed > 0:
                cache_ratio = (total_cached / (total_cached + total_computed)) * 100
                print(f"   - Cache efficiency: {cache_ratio:.1f}%")
            
            collection = vector_store.vector_store._collection
            count = collection.count()
            print(f"\nVector store now contains {count} document chunks")
            
            # Show sample
            if count > 0:
                sample = collection.get(limit=1, include=["metadatas", "documents"])
                print(f"\nSample chunk:")
                print(f"   Metadata: {sample['metadatas'][0]}")
                print(f"   Preview: {sample['documents'][0][:200]}...")
            
            return True
            
        except Exception as e:
            print(f"ERROR: Error building vector store: {e}")
            self.logger.error(f"Error building vector store: {e}")
            import traceback
            traceback.print_exc()
            return False


if __name__ == "__main__":
    processor = AdvancedDocumentProcessor(
        chunk_size=800,      # ~200 tokens
        chunk_overlap=120    # 15% overlap
    )
    
    documents = processor.process_directory()
    
    if documents:
        use_local = Config().USE_LOCAL_LLM
        processor.build_vector_store(documents, use_local=use_local)